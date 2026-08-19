import flet as ft
import os
import io
import base64
import threading
from PIL import Image

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def get_file_converter_control(page: ft.Page) -> ft.Control:
    """Returns the Flet Control UI for File Converter (Images, Documents, PDFs)."""

    status_text = ft.Text("", size=13, weight=ft.FontWeight.W_500)

    # Helper for status messages
    def set_status(msg: str, color=ft.Colors.GREEN_400):
        status_text.value = msg
        status_text.color = color
        page.update()

    # --- TAB 1: IMAGE CONVERTER & RESIZER ---
    selected_img_bytes = [None]
    selected_img_name = [""]

    img_preview = ft.Image(
        src="",
        width=180,
        height=180,
        fit="contain",
        visible=False,
        border_radius=12,
    )

    img_info_text = ft.Text("No image selected", size=13, color=ft.Colors.GREY_400)

    target_format_dropdown = ft.Dropdown(
        label="Target Format",
        value="PNG",
        options=[
            ft.dropdown.Option("PNG"),
            ft.dropdown.Option("JPEG"),
            ft.dropdown.Option("WEBP"),
            ft.dropdown.Option("BMP"),
            ft.dropdown.Option("PDF"),
        ],
        border_radius=10,
        expand=True,
    )

    quality_slider = ft.Slider(min=10, max=100, value=85, label="{value}% Quality")
    resize_scale = ft.Dropdown(
        label="Resize Scale",
        value="100%",
        options=[
            ft.dropdown.Option("100%"),
            ft.dropdown.Option("75%"),
            ft.dropdown.Option("50%"),
            ft.dropdown.Option("25%"),
        ],
        border_radius=10,
        expand=True,
    )

    converted_img_preview = ft.Image(
        src="",
        width=180,
        height=180,
        fit="contain",
        visible=False,
        border_radius=12,
    )
    download_btn = ft.Container(visible=False)

    def load_image_bytes(img_bytes: bytes, file_name: str):
        try:
            selected_img_bytes[0] = img_bytes
            selected_img_name[0] = file_name
            img = Image.open(io.BytesIO(img_bytes))
            img_info_text.value = f"Selected: {file_name} ({img.width}x{img.height} px, {img.format})"
            img_preview.src_base64 = base64.b64encode(img_bytes).decode("utf-8")
            img_preview.visible = True
            set_status("Image loaded successfully!", ft.Colors.GREEN_400)
        except Exception as ex:
            set_status(f"Error reading image: {ex}", ft.Colors.RED_400)
        page.update()

    def pick_image_file(_):
        def _open():
            try:
                from tkinter import Tk, filedialog
                root = Tk()
                root.withdraw()
                root.attributes('-topmost', True)
                fpath = filedialog.askopenfilename(
                    title="Select Image File",
                    filetypes=[
                        ("Image Files", "*.png *.jpg *.jpeg *.webp *.bmp"),
                        ("All Files", "*.*"),
                    ]
                )
                root.destroy()
                if fpath and os.path.exists(fpath):
                    with open(fpath, "rb") as f:
                        data = f.read()
                    load_image_bytes(data, os.path.basename(fpath))
            except Exception:
                # Fallback to Flet FilePicker if tkinter is unavailable
                def on_img_picked(e: ft.FilePickerResultEvent):
                    if e.files and len(e.files) > 0:
                        picked_file = e.files[0]
                        if picked_file.path and os.path.exists(picked_file.path):
                            with open(picked_file.path, "rb") as f:
                                b = f.read()
                            load_image_bytes(b, picked_file.name)
                        elif picked_file.bytes:
                            load_image_bytes(picked_file.bytes, picked_file.name)

                fp = ft.FilePicker()
                fp.on_result = on_img_picked
                page.overlay.append(fp)
                page.update()
                fp.pick_files(allow_multiple=False, file_type="image")

        threading.Thread(target=_open, daemon=True).start()

    def save_bytes_to_file(bytes_data: bytes, default_name: str):
        def _save():
            try:
                from tkinter import Tk, filedialog
                root = Tk()
                root.withdraw()
                root.attributes('-topmost', True)
                ext = os.path.splitext(default_name)[1]
                fpath = filedialog.asksaveasfilename(
                    title="Save Output File",
                    initialfile=default_name,
                    defaultextension=ext,
                    filetypes=[(f"*{ext}", f"*{ext}"), ("All Files", "*.*")]
                )
                root.destroy()
                if fpath:
                    with open(fpath, "wb") as f:
                        f.write(bytes_data)
                    set_status(f"File saved to {os.path.basename(fpath)}!", ft.Colors.GREEN_400)
            except Exception as ex:
                set_status(f"Save error: {ex}", ft.Colors.RED_400)

        threading.Thread(target=_save, daemon=True).start()

    def convert_image(_):
        if not selected_img_bytes[0]:
            set_status("Please select an image first!", ft.Colors.AMBER_400)
            return

        try:
            set_status("Converting image...", ft.Colors.BLUE_400)
            raw_img = Image.open(io.BytesIO(selected_img_bytes[0]))

            # Resize scale calculation
            scale = int(resize_scale.value.replace("%", "")) / 100.0
            new_w = max(1, int(raw_img.width * scale))
            new_h = max(1, int(raw_img.height * scale))
            res_img = raw_img.resize((new_w, new_h), Image.Resampling.LANCZOS)

            fmt = target_format_dropdown.value.upper()
            output_io = io.BytesIO()

            if fmt == "JPEG" or fmt == "JPG":
                if res_img.mode in ("RGBA", "P"):
                    res_img = res_img.convert("RGB")
                res_img.save(output_io, format="JPEG", quality=int(quality_slider.value))
                ext = "jpg"
            elif fmt == "PDF":
                if res_img.mode in ("RGBA", "P"):
                    res_img = res_img.convert("RGB")
                res_img.save(output_io, format="PDF")
                ext = "pdf"
            elif fmt == "WEBP":
                res_img.save(output_io, format="WEBP", quality=int(quality_slider.value))
                ext = "webp"
            elif fmt == "BMP":
                if res_img.mode in ("RGBA", "P"):
                    res_img = res_img.convert("RGB")
                res_img.save(output_io, format="BMP")
                ext = "bmp"
            else:  # PNG default
                res_img.save(output_io, format="PNG")
                ext = "png"

            out_bytes = output_io.getvalue()

            if fmt != "PDF":
                converted_img_preview.src_base64 = base64.b64encode(out_bytes).decode("utf-8")
                converted_img_preview.visible = True
            else:
                converted_img_preview.visible = False

            base_name = os.path.splitext(selected_img_name[0])[0] or "converted_file"
            file_out_name = f"{base_name}_converted.{ext}"

            download_btn.content = ft.Button(
                f"Save Converted File (.{ext.upper()})",
                icon=ft.Icons.DOWNLOAD_ROUNDED,
                on_click=lambda _: save_bytes_to_file(out_bytes, file_out_name),
                style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE),
            )
            download_btn.visible = True

            set_status(f"Success! Converted to {fmt} ({len(out_bytes) // 1024} KB)", ft.Colors.GREEN_400)
        except Exception as ex:
            set_status(f"Conversion error: {str(ex)}", ft.Colors.RED_400)
        page.update()

    image_tab_content = ft.Column([
        ft.Text("Image Converter & Resizer", size=16, weight=ft.FontWeight.BOLD),
        ft.Row([
            ft.Button(
                "Pick Image File",
                icon=ft.Icons.IMAGE_OUTLINED,
                on_click=pick_image_file,
            ),
        ]),
        img_info_text,
        img_preview,
        ft.Divider(),
        ft.Row([target_format_dropdown, resize_scale]),
        ft.Text("Quality / Compression:"),
        quality_slider,
        ft.Button("Convert Image", icon=ft.Icons.TRANSFORM, on_click=convert_image, style=ft.ButtonStyle(bgcolor=ft.Colors.INDIGO_600, color=ft.Colors.WHITE)),
        converted_img_preview,
        download_btn,
    ], spacing=12)

    # --- TAB 2: DOCUMENT & TEXT TO PDF CONVERTER ---
    doc_text_input = ft.TextField(
        hint_text="Paste or type text here to generate a Word (.docx) or PDF file...",
        multiline=True,
        min_lines=5,
        max_lines=10,
        border_radius=10,
    )
    doc_status = ft.Text("", size=13)

    def create_docx_file(_):
        txt = doc_text_input.value.strip()
        if not txt:
            doc_status.value = "Please enter some text first!"
            doc_status.color = ft.Colors.AMBER_400
            page.update()
            return

        try:
            out_io = io.BytesIO()
            if HAS_DOCX:
                doc = docx.Document()
                doc.add_heading("Generated Document", level=1)
                for line in txt.split("\n"):
                    doc.add_paragraph(line)
                doc.save(out_io)
                docx_bytes = out_io.getvalue()
            else:
                docx_bytes = txt.encode("utf-8")

            save_bytes_to_file(docx_bytes, "Document.docx")
            doc_status.value = "Preparing file download dialog..."
            doc_status.color = ft.Colors.GREEN_400
        except Exception as ex:
            doc_status.value = f"Error creating docx: {ex}"
            doc_status.color = ft.Colors.RED_400
        page.update()

    def create_pdf_from_text(_):
        txt = doc_text_input.value.strip()
        if not txt:
            doc_status.value = "Please enter some text first!"
            doc_status.color = ft.Colors.AMBER_400
            page.update()
            return

        try:
            lines = txt.split("\n")
            img = Image.new("RGB", (1240, 1754), color="white")
            from PIL import ImageDraw
            draw = ImageDraw.Draw(img)

            y = 100
            for line in lines:
                draw.text((80, y), line, fill="black")
                y += 40
                if y > 1650:
                    break

            pdf_io = io.BytesIO()
            img.save(pdf_io, format="PDF")
            pdf_bytes = pdf_io.getvalue()

            save_bytes_to_file(pdf_bytes, "Converted_Document.pdf")
            doc_status.value = "Preparing PDF download dialog..."
            doc_status.color = ft.Colors.GREEN_400
        except Exception as ex:
            doc_status.value = f"PDF generation error: {ex}"
            doc_status.color = ft.Colors.RED_400
        page.update()

    doc_tab_content = ft.Column([
        ft.Text("Document & Text Converter", size=16, weight=ft.FontWeight.BOLD),
        doc_text_input,
        doc_status,
        ft.Row([
            ft.Button(
                "Export as Word (.DOCX)",
                icon=ft.Icons.DESCRIPTION_ROUNDED,
                on_click=create_docx_file,
                style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_700, color=ft.Colors.WHITE),
            ),
            ft.Button(
                "Export as PDF (.PDF)",
                icon=ft.Icons.PICTURE_IN_PICTURE_ROUNDED,
                on_click=create_pdf_from_text,
                style=ft.ButtonStyle(bgcolor=ft.Colors.RED_700, color=ft.Colors.WHITE),
            ),
        ], wrap=True),
        ft.Divider(),
        ft.Text("Word (.docx) / File Features:", size=13, color=ft.Colors.GREY_400),
        ft.Text("• Convert images to PNG, JPG, WEBP, BMP & PDF\n• Scale and compress image file sizes\n• Export notes & text directly to Word (.docx) & PDF documents", size=12, color=ft.Colors.GREY_500),
    ], spacing=12)

    # --- SECTION SWITCHER CONTAINER ---
    content_area = ft.Container(content=image_tab_content, expand=True)

    btn_img_tab = ft.Button(
        "Image Converter 🖼️",
        style=ft.ButtonStyle(bgcolor=ft.Colors.INDIGO_700, color=ft.Colors.WHITE, shape=ft.RoundedRectangleBorder(radius=10)),
    )
    btn_doc_tab = ft.Button(
        "Doc & Text to PDF 📄",
        style=ft.ButtonStyle(bgcolor=ft.Colors.SURFACE_CONTAINER_HIGHEST, color=ft.Colors.WHITE, shape=ft.RoundedRectangleBorder(radius=10)),
    )

    def show_img_section(_):
        btn_img_tab.style.bgcolor = ft.Colors.INDIGO_700
        btn_doc_tab.style.bgcolor = ft.Colors.SURFACE_CONTAINER_HIGHEST
        content_area.content = image_tab_content
        page.update()

    def show_doc_section(_):
        btn_img_tab.style.bgcolor = ft.Colors.SURFACE_CONTAINER_HIGHEST
        btn_doc_tab.style.bgcolor = ft.Colors.INDIGO_700
        content_area.content = doc_tab_content
        page.update()

    btn_img_tab.on_click = show_img_section
    btn_doc_tab.on_click = show_doc_section

    toggle_bar = ft.Row([btn_img_tab, btn_doc_tab], alignment=ft.MainAxisAlignment.CENTER, wrap=True)

    return ft.Column([
        status_text,
        toggle_bar,
        ft.Divider(),
        content_area,
    ], expand=True, scroll=ft.ScrollMode.AUTO)
